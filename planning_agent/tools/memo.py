"""Memo tools - Word document generation for investment memos and system pitches."""

import os
from datetime import datetime
from pathlib import Path
from typing import Any, Optional

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

from planning_agent.client.planning_client import PlanningClient

_client: PlanningClient = None
_app_name: str = None


def set_client(client: PlanningClient):
    global _client
    _client = client


def set_app_name(app_name: str):
    global _app_name
    _app_name = app_name


def _add_memo_header(doc: Document, title: str, subtitle: str = "") -> None:
    """Add a styled header to the document."""
    # Title
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(18)
    title_run.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtitle
    if subtitle:
        subtitle_para = doc.add_paragraph()
        subtitle_run = subtitle_para.add_run(subtitle)
        subtitle_run.font.size = Pt(12)
        subtitle_run.font.color.rgb = RGBColor(102, 102, 102)  # Gray
        subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Date
    date_para = doc.add_paragraph()
    date_run = date_para.add_run(datetime.now().strftime("%B %d, %Y"))
    date_run.font.size = Pt(10)
    date_run.font.color.rgb = RGBColor(128, 128, 128)
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # Spacing


def _add_section_header(doc: Document, text: str) -> None:
    """Add a section header."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 51, 102)


def _add_bullet_list(doc: Document, items: list[str]) -> None:
    """Add a bulleted list."""
    for item in items:
        para = doc.add_paragraph(item, style='List Bullet')


def _add_key_value_table(doc: Document, data: dict[str, Any]) -> None:
    """Add a simple key-value table."""
    table = doc.add_table(rows=len(data), cols=2)
    table.style = 'Table Grid'

    for i, (key, value) in enumerate(data.items()):
        row = table.rows[i]
        row.cells[0].text = str(key)
        row.cells[1].text = str(value)

    doc.add_paragraph()  # Spacing


async def _get_financial_data(
    entity: str = "E501",
    years: str = "FY25",
    scenario: str = "Actual"
) -> dict[str, Any]:
    """Get financial data for memo generation."""
    from planning_agent.tools import data

    # Get revenue data
    revenue_result = await data.smart_retrieve_revenue(
        entity=entity,
        years=years,
        scenario=scenario,
        period="YearTotal"
    )

    # Get variance data for key accounts
    variance_result = await data.smart_retrieve_variance(
        account="400000",  # Total Revenue
        entity=entity,
        years=years
    )

    return {
        "revenue": revenue_result.get("data", {}),
        "variance": variance_result.get("data", {})
    }


def _analyze_financials(data: dict[str, Any]) -> dict[str, Any]:
    """Analyze financial data for insights."""
    revenue_data = data.get("revenue", {})
    variance_data = data.get("variance", {})

    summary = revenue_data.get("summary", {})
    variance = variance_data.get("variance", {})

    total_revenue = summary.get("total_revenue_400000", 0)
    rooms_revenue = summary.get("rooms_revenue_410000", 0)
    fb_revenue = summary.get("fb_revenue_420000", 0)

    # Calculate mix
    rooms_mix = (rooms_revenue / total_revenue * 100) if total_revenue else 0
    fb_mix = (fb_revenue / total_revenue * 100) if total_revenue else 0

    return {
        "total_revenue": total_revenue,
        "rooms_revenue": rooms_revenue,
        "fb_revenue": fb_revenue,
        "rooms_mix_pct": round(rooms_mix, 1),
        "fb_mix_pct": round(fb_mix, 1),
        "yoy_growth_pct": variance.get("actual_vs_prior_year_pct", 0),
        "forecast_variance_pct": variance.get("actual_vs_forecast_pct", 0)
    }


async def generate_system_pitch(
    output_path: Optional[str] = None
) -> dict[str, Any]:
    """Generate a 1-page system overview Word document.

    Creates a professional document explaining the Oracle EPM Planning
    AI Assistant capabilities, features, and benefits.

    Args:
        output_path: Optional path for output file (default: ./output/system_pitch.docx)

    Returns:
        dict: Status and path to generated document.
    """
    try:
        doc = Document()

        # Header
        _add_memo_header(
            doc,
            "Oracle EPM Planning AI Assistant",
            "Intelligent Financial Planning & Analysis"
        )

        # Executive Summary
        _add_section_header(doc, "Executive Summary")
        doc.add_paragraph(
            "The Oracle EPM Planning AI Assistant is an intelligent agent that enables "
            "natural language interaction with Oracle EPM Cloud Planning applications. "
            "It combines advanced AI capabilities with deep Oracle EPM expertise to "
            "streamline financial planning, analysis, and reporting workflows."
        )

        # Key Capabilities
        _add_section_header(doc, "Key Capabilities")
        _add_bullet_list(doc, [
            "Natural Language Queries: Ask questions in plain English or Portuguese",
            "Smart Data Retrieval: Automatic dimension handling for complex cube structures",
            "Variance Analysis: Instant comparisons across scenarios, periods, and years",
            "Intelligent Member Resolution: Fuzzy matching and semantic search for dimensions",
            "Document Generation: Automated investment memos and financial reports",
            "Reinforcement Learning: Continuously improving based on user feedback"
        ])

        # Supported Operations
        _add_section_header(doc, "Supported Operations")
        _add_bullet_list(doc, [
            "Query financial data across all dimensions (Account, Entity, Period, etc.)",
            "Execute business rules and calculations",
            "Monitor and manage Planning jobs",
            "Explore dimension hierarchies and members",
            "Analyze variances (Actual vs Forecast, YoY comparisons)",
            "Generate professional Word documents and reports"
        ])

        # Technical Architecture
        _add_section_header(doc, "Technical Architecture")
        _add_key_value_table(doc, {
            "Integration": "Oracle EPM REST API v3",
            "AI Model": "Claude (Anthropic)",
            "Protocol": "Model Context Protocol (MCP)",
            "Database": "SQLite for sessions, preferences, and RL",
            "Languages": "Python, TypeScript"
        })

        # Benefits
        _add_section_header(doc, "Benefits")
        _add_bullet_list(doc, [
            "Reduce time spent on routine data retrieval by 80%",
            "Enable non-technical users to access Planning data",
            "Improve accuracy with intelligent validation",
            "Accelerate financial close and reporting cycles",
            "Learn from usage patterns to improve recommendations"
        ])

        # Save document
        if output_path is None:
            output_dir = Path("./output")
            output_dir.mkdir(exist_ok=True)
            output_path = str(output_dir / "system_pitch.docx")

        doc.save(output_path)

        return {
            "status": "success",
            "message": "System pitch document generated successfully",
            "data": {
                "path": output_path,
                "pages": 1,
                "generated_at": datetime.now().isoformat()
            }
        }

    except Exception as e:
        return {"status": "error", "error": str(e)}


async def generate_investment_memo(
    entity: str = "E501",
    entity_name: str = "L7 Chicago",
    years: str = "FY25",
    output_path: Optional[str] = None
) -> dict[str, Any]:
    """Generate a 2-page investment memo Word document.

    Creates a professional financial analysis memo for the specified entity,
    including revenue breakdown, variance analysis, and recommendations.

    Args:
        entity: Entity code (default: 'E501')
        entity_name: Entity display name (default: 'L7 Chicago')
        years: Fiscal year (default: 'FY25')
        output_path: Optional path for output file

    Returns:
        dict: Status and path to generated document.
    """
    try:
        # Get financial data
        financial_data = await _get_financial_data(entity=entity, years=years)
        analysis = _analyze_financials(financial_data)

        doc = Document()

        # Header
        _add_memo_header(
            doc,
            f"Investment Memo: {entity_name}",
            f"Financial Analysis - {years}"
        )

        # Executive Summary
        _add_section_header(doc, "Executive Summary")

        total_rev = analysis["total_revenue"]
        yoy_growth = analysis["yoy_growth_pct"]

        summary_text = (
            f"{entity_name} ({entity}) generated ${total_rev:,.0f} in total revenue "
            f"for {years}, representing a {yoy_growth:+.1f}% change versus prior year. "
        )

        if yoy_growth > 5:
            summary_text += "The entity demonstrates strong growth momentum. "
        elif yoy_growth < -5:
            summary_text += "Performance challenges warrant attention. "
        else:
            summary_text += "Performance is tracking in line with expectations. "

        doc.add_paragraph(summary_text)

        # Financial Highlights
        _add_section_header(doc, "Financial Highlights")
        _add_key_value_table(doc, {
            "Total Revenue": f"${analysis['total_revenue']:,.0f}",
            "Rooms Revenue": f"${analysis['rooms_revenue']:,.0f} ({analysis['rooms_mix_pct']}% mix)",
            "F&B Revenue": f"${analysis['fb_revenue']:,.0f} ({analysis['fb_mix_pct']}% mix)",
            "YoY Growth": f"{analysis['yoy_growth_pct']:+.1f}%",
            "vs Forecast": f"{analysis['forecast_variance_pct']:+.1f}%"
        })

        # Revenue Analysis
        _add_section_header(doc, "Revenue Analysis")

        rooms_commentary = (
            f"Rooms revenue of ${analysis['rooms_revenue']:,.0f} represents "
            f"{analysis['rooms_mix_pct']:.1f}% of total revenue. "
        )
        if analysis['rooms_mix_pct'] > 60:
            rooms_commentary += "The strong rooms contribution reflects solid occupancy rates."
        else:
            rooms_commentary += "Opportunity exists to drive higher rooms contribution."

        doc.add_paragraph(rooms_commentary)

        fb_commentary = (
            f"Food & Beverage revenue of ${analysis['fb_revenue']:,.0f} represents "
            f"{analysis['fb_mix_pct']:.1f}% of total revenue. "
        )
        doc.add_paragraph(fb_commentary)

        # Variance Analysis
        _add_section_header(doc, "Variance Analysis")

        var_text = f"Year-over-year performance shows a {analysis['yoy_growth_pct']:+.1f}% change. "
        if analysis['forecast_variance_pct'] > 0:
            var_text += f"Actual results exceeded forecast by {analysis['forecast_variance_pct']:.1f}%."
        elif analysis['forecast_variance_pct'] < 0:
            var_text += f"Actual results fell short of forecast by {abs(analysis['forecast_variance_pct']):.1f}%."
        else:
            var_text += "Actual results are in line with forecast."

        doc.add_paragraph(var_text)

        # Recommendations
        _add_section_header(doc, "Recommendations")

        recommendations = []
        if analysis['yoy_growth_pct'] < 0:
            recommendations.append("Review pricing strategy and competitive positioning")
            recommendations.append("Analyze root causes of revenue decline")
        if analysis['forecast_variance_pct'] < -5:
            recommendations.append("Update forecast assumptions based on current trends")
            recommendations.append("Implement corrective actions to close the gap")
        if analysis['rooms_mix_pct'] < 50:
            recommendations.append("Focus on rooms revenue optimization")
        if analysis['fb_mix_pct'] < 20:
            recommendations.append("Explore F&B enhancement opportunities")

        if not recommendations:
            recommendations.append("Continue executing current strategy")
            recommendations.append("Monitor key performance indicators monthly")

        _add_bullet_list(doc, recommendations)

        # Page break and next steps
        doc.add_page_break()

        _add_section_header(doc, "Next Steps")
        _add_bullet_list(doc, [
            "Schedule deep-dive review with property management",
            "Analyze monthly trends for seasonal patterns",
            "Compare performance against peer group",
            "Review operating expense efficiency"
        ])

        # Appendix - Data Sources
        _add_section_header(doc, "Appendix: Data Sources")
        doc.add_paragraph(
            f"Data retrieved from Oracle EPM Planning ({_app_name or 'PlanApp'}) "
            f"as of {datetime.now().strftime('%Y-%m-%d %H:%M')}. "
            "Revenue figures represent actual results for the period indicated."
        )

        # Save document
        if output_path is None:
            output_dir = Path("./output")
            output_dir.mkdir(exist_ok=True)
            output_path = str(output_dir / f"investment_memo_{entity}_{years}.docx")

        doc.save(output_path)

        return {
            "status": "success",
            "message": f"Investment memo for {entity_name} generated successfully",
            "data": {
                "path": output_path,
                "entity": entity,
                "entity_name": entity_name,
                "years": years,
                "pages": 2,
                "generated_at": datetime.now().isoformat(),
                "highlights": analysis
            }
        }

    except Exception as e:
        return {"status": "error", "error": str(e)}


TOOL_DEFINITIONS = [
    {
        "name": "generate_system_pitch",
        "description": "Generate a 1-page Word document describing the Oracle EPM Planning AI Assistant capabilities / Gerar documento de apresentacao do sistema",
        "inputSchema": {
            "type": "object",
            "properties": {
                "output_path": {
                    "type": "string",
                    "description": "Optional path for output file (default: ./output/system_pitch.docx)",
                },
            },
        },
    },
    {
        "name": "generate_investment_memo",
        "description": "Generate a 2-page investment memo Word document with financial analysis for an entity / Gerar memo de investimento com analise financeira",
        "inputSchema": {
            "type": "object",
            "properties": {
                "entity": {
                    "type": "string",
                    "description": "Entity code (default: 'E501')",
                },
                "entity_name": {
                    "type": "string",
                    "description": "Entity display name (default: 'L7 Chicago')",
                },
                "years": {
                    "type": "string",
                    "description": "Fiscal year (default: 'FY25')",
                },
                "output_path": {
                    "type": "string",
                    "description": "Optional path for output file",
                },
            },
        },
    },
]
